"""공직 채용정보(PblJobService) MCP 서버.

인사혁신처 공직 채용정보 조회 서비스(apis.data.go.kr)를 MCP 도구 3개로 노출한다.
- search_jobs: 채용공고 목록 검색 (getList)
- get_job_detail: 공고 상세 조회 (getItem)
- get_job_files: 공고 첨부파일 조회 (getItemFile)

인증키는 환경변수 PUBJOB_API_KEY에서 읽는다 (.env 파일 지원).
"""

import asyncio
import json
import os
import re
import tempfile
import zipfile
import zlib
from datetime import date, datetime, timedelta
from typing import Any
from urllib.parse import quote, unquote

import httpx
import xmltodict
from dotenv import load_dotenv
from mcp.server import MCPServer

load_dotenv()

BASE_URL = "https://apis.data.go.kr/1760000/PblJobService"
TIMEOUT = 10.0

INSTT_SE = {
    "g01": "국가공무원",
    "g02": "지방공무원",
    "g03": "공공기관",
    "g04": "교육청",
}
PBLANC_TY = {
    "e01": "공개경쟁",
    "e02": "경력경쟁",
    "e03": "계약직",
    "e04": "행정지원",
    "e06": "공모직위",
}

mcp = MCPServer(
    "pubjob-mcp",
    instructions="대한민국 공직(국가·지방공무원, 공공기관, 교육청) 채용공고를 검색·조회하는 서버",
)


async def _call_api(path: str, params: dict[str, Any]) -> dict[str, Any] | str:
    """API를 호출해 응답 body(dict)를 반환한다. 실패 시 에러 메시지(str)를 반환한다."""
    key = os.environ.get("PUBJOB_API_KEY")
    if not key:
        return "오류: 환경변수 PUBJOB_API_KEY가 비어 있습니다. .env 파일에 인증키를 입력하세요."
    if "%" in key:  # Encoding 키가 들어온 경우 원본(Decoding)으로 되돌린다 (httpx가 다시 인코딩하므로)
        key = unquote(key)

    query: dict[str, Any] = {"serviceKey": key, "pageNo": "1", "numOfRows": "10"}
    query.update({k: v for k, v in params.items() if v not in (None, "")})

    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            resp = await client.get(f"{BASE_URL}/{path}", params=query)
            resp.raise_for_status()
    except httpx.TimeoutException:
        return "오류: API 응답이 10초를 초과했습니다. 잠시 후 다시 시도하세요."
    except httpx.HTTPError as e:
        return f"오류: API 요청 실패 - {e}"

    try:
        parsed = xmltodict.parse(resp.text)
    except Exception as e:
        return f"오류: XML 파싱 실패 - {e} / 응답 일부: {resp.text[:200]}"

    # data.go.kr 게이트웨이 공통 에러(인증키 오류 등)는 response가 아니라 이 형식으로 온다
    if "OpenAPI_ServiceResponse" in parsed:
        hdr = (parsed["OpenAPI_ServiceResponse"] or {}).get("cmmMsgHeader") or {}
        reason = hdr.get("returnAuthMsg") or hdr.get("errMsg") or "설명 없음"
        return f"오류: API 게이트웨이 에러 (returnReasonCode={hdr.get('returnReasonCode')}, {reason})"

    response = parsed.get("response") or {}
    header = response.get("header") or {}
    code = header.get("resultCode")
    if code != "00":
        return f"오류: API 실패 응답 (resultCode={code}, resultMsg={header.get('resultMsg', '설명 없음')})"
    return response.get("body") or {}


def _extract_items(body: dict[str, Any]) -> list[dict[str, Any]]:
    """body에서 item을 꺼내 1건(dict)이든 여러 건(list)이든 list로 통일한다."""
    items = body.get("items")
    if isinstance(items, dict) and "item" in items:
        items = items["item"]
    if items is None:
        items = body.get("item")
    if items is None:
        return []
    if isinstance(items, list):
        return [i for i in items if isinstance(i, dict)]
    if isinstance(items, dict):
        return [items]
    return []


def _fmt_value(value: Any) -> str:
    if value is None:
        return "-"
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


@mcp.tool()
async def search_jobs(
    keyword: str | None = None,
    instt_se: str | None = None,
    instt_nm: str | None = None,
    pblanc_ty: str | None = None,
    begin_de: str | None = None,
    end_de: str | None = None,
    num_rows: int = 10,
) -> str:
    """공직 채용공고를 검색한다. 모든 인자는 선택적이며 항상 최신 등록순으로 반환한다.

    Args:
        keyword: 제목 검색 키워드 (예: "전산", "간호")
        instt_se: 기관 구분 코드 — g01=국가공무원, g02=지방공무원, g03=공공기관, g04=교육청
        instt_nm: 기관명 (예: "서울특별시", "국세청")
        pblanc_ty: 공고 유형 코드 — e01=공개경쟁, e02=경력경쟁, e03=계약직, e04=행정지원, e06=공모직위
        begin_de: 조회 시작일 YYYYMMDD (미지정 시 오늘-30일)
        end_de: 조회 종료일 YYYYMMDD (미지정 시 오늘)
        num_rows: 최대 결과 수 (기본 10)
    """
    today = date.today()
    begin_de = begin_de or (today - timedelta(days=30)).strftime("%Y%m%d")
    end_de = end_de or today.strftime("%Y%m%d")

    body = await _call_api(
        "getList",
        {
            "Kwrd": keyword,
            "Instt_se": instt_se,
            "Instt_nm": instt_nm,
            "Pblanc_ty": pblanc_ty,
            "Begin_de": begin_de,
            "End_de": end_de,
            "Sort_order": "DESC",  # 항상 내림차순(최신순) 고정
            "numOfRows": str(max(1, num_rows)),
        },
    )
    if isinstance(body, str):
        return body

    conditions = [f"기간: {begin_de}~{end_de}"]
    if keyword:
        conditions.append(f"키워드: {keyword}")
    if instt_se:
        conditions.append(f"기관구분: {INSTT_SE.get(instt_se, instt_se)}({instt_se})")
    if instt_nm:
        conditions.append(f"기관명: {instt_nm}")
    if pblanc_ty:
        conditions.append(f"공고유형: {PBLANC_TY.get(pblanc_ty, pblanc_ty)}({pblanc_ty})")
    cond_text = " / ".join(conditions)

    items = _extract_items(body)
    if not items:
        return f"검색 결과 없음\n[검색 조건] {cond_text}"

    total = body.get("totalCount")
    lines = [f"[검색 조건] {cond_text}"]
    if total is not None:
        lines.append(f"전체 {total}건 중 {len(items)}건 표시")
    for it in items:
        lines.append(
            f"- [{_fmt_value(it.get('idx'))}] {_fmt_value(it.get('title'))}\n"
            f"  기관: {_fmt_value(it.get('insttname'))} | "
            f"등록일: {_fmt_value(it.get('regdate'))} | "
            f"마감일: {_fmt_value(it.get('enddate'))}"
        )
    return "\n".join(lines)


@mcp.tool()
async def get_job_detail(idx: str) -> str:
    """채용공고 1건의 상세 정보를 조회한다.

    Args:
        idx: 공고 식별 번호 (search_jobs 결과의 [idx] 값)
    """
    body = await _call_api("getItem", {"idx": idx})
    if isinstance(body, str):
        return body

    items = _extract_items(body)
    if not items:
        return f"공고를 찾을 수 없습니다 (idx={idx})"

    lines = [f"[공고 상세] idx={idx}"]
    for item in items:
        for key, value in item.items():
            lines.append(f"{key}: {_fmt_value(value)}")
    return "\n".join(lines)


@mcp.tool()
async def get_job_files(idx: str) -> str:
    """채용공고의 첨부파일 목록(파일명·다운로드 정보)을 조회한다.

    Args:
        idx: 공고 식별 번호 (search_jobs 결과의 [idx] 값)
    """
    body = await _call_api("getItemFile", {"idx": idx})
    if isinstance(body, str):
        return body

    items = _extract_items(body)
    if not items:
        return f"첨부파일 없음 (idx={idx})"

    lines = [f"[첨부파일] idx={idx} / {len(items)}건"]
    for i, item in enumerate(items, 1):
        lines.append(f"({i})")
        for key, value in item.items():
            lines.append(f"  {key}: {_fmt_value(value)}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# 잡알리오 (공공기관 채용공시, opendata.alio.go.kr)
# ---------------------------------------------------------------------------

ALIO_BASE_URL = "https://opendata.alio.go.kr"

ALIO_ERROR_CODES = {
    "3": "데이터 없음",
    "7": "게이트웨이 인증 실패",
    "8": "활용신청 승인 확인 실패",
}

ALIO_FILE_TYPES = {"A": "공고문", "B": "입사지원서", "C": "직무기술서"}


async def _call_alio(path: str, params: dict[str, Any]) -> dict[str, Any] | str:
    """잡알리오 API 호출. 정상 시 응답 JSON(dict), 실패 시 에러 메시지(str) 반환.

    명세상 POST이지만 파라미터는 전부 query string으로 보낸다.
    POST가 실패하면 GET으로 재시도한다.
    """
    key = os.environ.get("ALIO_API_KEY")
    if not key:
        return "오류: 환경변수 ALIO_API_KEY가 비어 있습니다. .env 파일에 인증키를 입력하세요."
    if "%" in key:  # Encoding 키가 들어온 경우 원본(Decoding)으로 되돌린다
        key = unquote(key)

    query: dict[str, Any] = {"serviceKey": key}
    query.update({k: v for k, v in params.items() if v not in (None, "")})

    url = f"{ALIO_BASE_URL}{path}"
    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            try:
                resp = await client.post(url, params=query)
                resp.raise_for_status()
            except httpx.HTTPError:
                resp = await client.get(url, params=query)  # POST 실패 시 GET 재시도
                resp.raise_for_status()
    except httpx.TimeoutException:
        return "오류: 잡알리오 API 응답이 10초를 초과했습니다. 잠시 후 다시 시도하세요."
    except httpx.HTTPError as e:
        return f"오류: 잡알리오 API 요청 실패 - {e}"

    try:
        data = resp.json()
    except Exception as e:
        return f"오류: JSON 파싱 실패 - {e} / 응답 일부: {resp.text[:200]}"

    code = str(data.get("resultCode"))
    if code not in ("0", "200"):  # 명세는 0이지만 실제 서버는 성공 시 200을 반환한다
        name = ALIO_ERROR_CODES.get(code, "")
        name_part = f" {name}," if name else ""
        return (
            f"오류: 잡알리오 API 실패 응답 (resultCode={code},{name_part} "
            f"resultMsg={data.get('resultMsg', '설명 없음')})"
        )
    return data


def _alio_items(data: dict[str, Any]) -> list[dict[str, Any]]:
    """응답에서 결과 목록을 꺼내 dict 1건이든 list든 list로 통일한다."""
    result = data.get("result")
    if result is None:
        result = data.get("items")
    if result is None:
        return []
    if isinstance(result, dict):
        return [result]
    if isinstance(result, list):
        return [r for r in result if isinstance(r, dict)]
    return []


@mcp.tool()
async def search_alio_jobs(
    keyword: str | None = None,
    ongoing_only: bool = True,
    hire_type: str | None = None,
    ncs_code: str | None = None,
    recrut_se: str | None = None,
    inst_type: str | None = None,
    inst_clsf: str | None = None,
    region: str | None = None,
    begin_date: str | None = None,
    end_date: str | None = None,
    num_rows: int = 20,
    page_no: int = 1,
) -> str:
    """잡알리오에서 공공기관 채용공시를 검색한다. 모든 인자는 선택적.

    Args:
        keyword: 공고 제목 포함검색 키워드 (서버측 검색)
        ongoing_only: True(기본)면 접수 진행 중인 공고만
        hire_type: 고용형태 코드(쉼표로 복수 지정) — R1010=정규직, R1020=계약직,
            R1030=무기계약직, R1040=비정규직, R1050=청년인턴,
            R1060=청년인턴(체험형), R1070=청년인턴(채용형)
        ncs_code: NCS 직무분야 코드(쉼표로 복수 지정) — R600001=사업관리,
            R600002=경영회계사무, R600004=교육자연사회과학, R600006=보건의료,
            R600007=사회복지종교, R600020=정보통신, R600025=연구
        recrut_se: 채용구분 코드 — R2010=신입, R2020=경력, R2030=신입+경력,
            R2040=외국인전형
        inst_type: 기관유형 코드 — A2001=공기업(시장형), A2002=공기업(준시장형),
            A2003=준정부(기금관리), A2004=준정부(위탁집행), A2005=기타공공기관
        inst_clsf: 기관분류 코드 — 04=고용보건복지, 08=연구교육
        region: 근무지역 코드(쉼표로 복수 지정) — R3010=서울, R3017=경기, R3026=세종 등
        begin_date: 공고 시작일 YYYY-MM-DD
        end_date: 공고 종료일 YYYY-MM-DD
        num_rows: 페이지당 결과 수 (기본 20)
        page_no: 페이지 번호 (기본 1)
    """
    data = await _call_alio(
        "/new/v1/recruit/list.do",
        {
            # 잡알리오 서버가 쿼리를 두 번 디코딩하므로 한글 키워드는 미리 한 번 인코딩한다
            # (미리 안 하면 resultCode=6 서버 에러. ASCII 키워드는 인코딩해도 동일)
            "recrutPbancTtl": quote(keyword) if keyword else None,
            "ongoingYn": "Y" if ongoing_only else None,
            "hireTypeLst": hire_type,
            "ncsCdLst": ncs_code,
            "recrutSe": recrut_se,
            "instType": inst_type,
            "instClsf": inst_clsf,
            "workRgnLst": region,
            "pbancBgngYmd": begin_date,
            "pbancEndYmd": end_date,
            "numOfRows": str(max(1, num_rows)),
            "pageNo": str(max(1, page_no)),
        },
    )
    if isinstance(data, str):
        return data

    conditions = []
    if keyword:
        conditions.append(f"키워드: {keyword}")
    if ongoing_only:
        conditions.append("진행 중인 공고만")
    if hire_type:
        conditions.append(f"고용형태: {hire_type}")
    if ncs_code:
        conditions.append(f"NCS: {ncs_code}")
    if recrut_se:
        conditions.append(f"채용구분: {recrut_se}")
    if inst_type:
        conditions.append(f"기관유형: {inst_type}")
    if inst_clsf:
        conditions.append(f"기관분류: {inst_clsf}")
    if region:
        conditions.append(f"지역: {region}")
    if begin_date or end_date:
        conditions.append(f"기간: {begin_date or ''}~{end_date or ''}")
    cond_text = " / ".join(conditions) if conditions else "조건 없음"

    items = _alio_items(data)
    if not items:
        return f"검색 결과 없음\n[검색 조건] {cond_text}"

    total = data.get("totalCount")
    lines = [f"[검색 조건] {cond_text}"]
    if total is not None:
        lines.append(f"전체 {total}건 중 {len(items)}건 표시 (페이지 {page_no})")
    for it in items:
        lines.append(
            f"- [{_fmt_value(it.get('recrutPblntSn'))}] {_fmt_value(it.get('recrutPbancTtl'))}\n"
            f"  기관: {_fmt_value(it.get('instNm'))} | "
            f"고용형태: {_fmt_value(it.get('hireTypeNmLst'))} | "
            f"채용구분: {_fmt_value(it.get('recrutSeNm'))}\n"
            f"  공고기간: {_fmt_value(it.get('pbancBgngYmd'))}~{_fmt_value(it.get('pbancEndYmd'))} | "
            f"지역: {_fmt_value(it.get('workRgnNmLst'))} | "
            f"NCS: {_fmt_value(it.get('ncsCdNmLst'))}"
        )
    return "\n".join(lines)


@mcp.tool()
async def get_alio_job_detail(sn: int) -> str:
    """잡알리오 채용공시 1건의 상세 정보(자격요건·우대·전형절차·첨부파일 등)를 조회한다.

    Args:
        sn: 공고 일련번호 (search_alio_jobs 결과의 [sn] 값, recrutPblntSn)
    """
    data = await _call_alio("/new/v1/recruit/detail.do", {"sn": sn})
    if isinstance(data, str):
        return data

    items = _alio_items(data)
    if not items:
        return f"공고를 찾을 수 없습니다 (sn={sn})"
    item = items[0]

    lines = [f"[공고 상세] sn={sn}"]
    lines.append(f"제목: {_fmt_value(item.get('recrutPbancTtl'))}")
    lines.append(f"기관: {_fmt_value(item.get('instNm'))}")
    lines.append(
        f"공고기간: {_fmt_value(item.get('pbancBgngYmd'))}~{_fmt_value(item.get('pbancEndYmd'))}"
    )
    for label, field in [
        ("고용형태", "hireTypeNmLst"),
        ("채용구분", "recrutSeNm"),
        ("근무지역", "workRgnNmLst"),
        ("NCS", "ncsCdNmLst"),
    ]:
        if item.get(field) is not None:
            lines.append(f"{label}: {_fmt_value(item.get(field))}")

    for label, field in [
        ("자격요건", "aplyQlfcCn"),
        ("우대사항", "prefCn"),
        ("우대조건", "prefCondCn"),
        ("결격사유", "disqlfcRsn"),
        ("전형절차", "scrnprcdrMthdExpln"),
    ]:
        value = item.get(field)
        if value not in (None, ""):
            lines.append(f"--- {label} ---")
            lines.append(_fmt_value(value))

    steps = item.get("steps")
    if isinstance(steps, dict):
        steps = [steps]
    if isinstance(steps, list) and steps:
        lines.append("--- 전형단계 ---")
        for i, step in enumerate(steps, 1):
            if not isinstance(step, dict):
                continue
            parts = []
            for k, v in step.items():
                if k == "aplyNope":
                    parts.append(f"지원인원: {_fmt_value(v)}")
                elif k == "cmpttRt":
                    parts.append(f"경쟁률: {_fmt_value(v)}")
                else:
                    parts.append(f"{k}: {_fmt_value(v)}")
            lines.append(f"({i}) " + " | ".join(parts))

    files = item.get("files")
    if isinstance(files, dict):
        files = [files]
    if isinstance(files, list) and files:
        lines.append("--- 첨부파일 ---")
        for f in files:
            if not isinstance(f, dict):
                continue
            ftype = ALIO_FILE_TYPES.get(str(f.get("atchFileType")), _fmt_value(f.get("atchFileType")))
            lines.append(f"- [{ftype}] {_fmt_value(f.get('atchFileNm'))} | {_fmt_value(f.get('url'))}")

    if item.get("srcUrl"):
        lines.append(f"원문링크: {item['srcUrl']}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# 첨부파일 텍스트 추출
# ---------------------------------------------------------------------------

DOC_TIMEOUT = 30.0
DOC_MAX_BYTES = 10 * 1024 * 1024  # 10MB
DOC_MAX_CHARS = 30_000

# HWP 5.0 인라인 컨트롤 중 8워드(16바이트)를 차지하는 확장 컨트롤 코드
_HWP_EXTENDED_CTRL = {1, 2, 3, 11, 12, 14, 15, 16, 17, 18, 21, 22, 23}


def _normalize_doc_url(url: str) -> str:
    """알리오 opendata 다운로드 URL은 파일 대신 HTML 페이지를 반환하므로
    실제 파일을 주는 www.alio.go.kr 패턴으로 변환한다."""
    m = re.match(
        r"https?://opendata\.alio\.go\.kr/recruit/downloadAtchFile\?recrutAtchFileNo=(\d+)",
        url,
    )
    if m:
        return f"https://www.alio.go.kr/download/download.json?fileNo={m.group(1)}"
    return url


async def _download_doc(url: str) -> tuple[str, str]:
    """파일을 임시폴더에 내려받아 (경로, 파일명 힌트)를 반환한다. 10MB 상한."""
    async with httpx.AsyncClient(timeout=DOC_TIMEOUT, follow_redirects=True) as client:
        async with client.stream("GET", url) as resp:
            resp.raise_for_status()
            length = resp.headers.get("content-length")
            if length and int(length) > DOC_MAX_BYTES:
                raise ValueError(f"파일이 10MB를 초과합니다 ({int(length):,} bytes)")

            # Content-Disposition에서 파일명 힌트 추출 (확장자 판별 보조용)
            filename = ""
            cd = resp.headers.get("content-disposition", "")
            m = re.search(r"filename\*=(?:UTF-8''|utf-8'')([^;]+)", cd)
            if not m:
                # 알리오는 filename=""이름""; 처럼 따옴표를 겹쳐 보내는 경우가 있다
                m = re.search(r'filename="*([^";]+)"*', cd)
            if m:
                filename = unquote(m.group(1).strip())

            chunks: list[bytes] = []
            total = 0
            async for chunk in resp.aiter_bytes():
                total += len(chunk)
                if total > DOC_MAX_BYTES:
                    raise ValueError("파일이 10MB를 초과합니다 (수신 중단)")
                chunks.append(chunk)

    fd, path = tempfile.mkstemp(prefix="pubjob_doc_")
    with os.fdopen(fd, "wb") as f:
        f.write(b"".join(chunks))
    return path, filename


def _detect_doc_format(path: str, filename_hint: str) -> str | None:
    """매직 바이트 우선, 파일명 확장자 보조로 문서 형식을 판별한다."""
    with open(path, "rb") as f:
        head = f.read(8)
    if head.startswith(b"%PDF"):
        return "pdf"
    if head.startswith(b"\xd0\xcf\x11\xe0"):  # OLE 컨테이너 (hwp/doc 공용)
        return "hwp"
    if head.startswith(b"PK"):  # ZIP 컨테이너 (hwpx/docx 공용) — 내용물로 구분
        try:
            with zipfile.ZipFile(path) as zf:
                names = zf.namelist()
        except zipfile.BadZipFile:
            return None
        if any(n.startswith("word/") for n in names):
            return "docx"
        if any(n.startswith("Contents/section") for n in names):
            return "hwpx"
        return None
    ext = os.path.splitext(filename_hint)[1].lower().lstrip(".")
    return ext if ext in ("hwp", "hwpx", "pdf", "docx") else None


def _extract_hwp(path: str) -> str:
    """HWP 5.0(OLE)의 BodyText 섹션을 zlib 해제 후 PARA_TEXT 레코드에서 텍스트 추출.

    표 셀 문단도 개별 PARA_TEXT 레코드로 저장되므로 표 안 텍스트가 함께 추출된다.
    """
    import olefile

    ole = olefile.OleFileIO(path)
    try:
        if not ole.exists("FileHeader"):
            raise ValueError("FileHeader 스트림이 없어 HWP 5.0 형식이 아닙니다")
        flags = int.from_bytes(ole.openstream("FileHeader").read()[36:40], "little")
        if flags & 0x02:
            raise ValueError("암호화(배포용) HWP 문서는 지원하지 않습니다")
        compressed = bool(flags & 0x01)

        sections = sorted(
            (e for e in ole.listdir() if e[0] == "BodyText"),
            key=lambda e: int(e[1].replace("Section", "")),
        )
        if not sections:
            raise ValueError("BodyText 섹션이 없습니다")

        out: list[str] = []
        for entry in sections:
            data = ole.openstream(entry).read()
            if compressed:
                data = zlib.decompress(data, -15)
            i = 0
            while i + 4 <= len(data):
                rec = int.from_bytes(data[i : i + 4], "little")
                tag = rec & 0x3FF
                size = (rec >> 20) & 0xFFF
                i += 4
                if size == 0xFFF:  # 확장 크기
                    size = int.from_bytes(data[i : i + 4], "little")
                    i += 4
                if tag == 67:  # HWPTAG_PARA_TEXT
                    chunk = data[i : i + size]
                    j = 0
                    buf: list[str] = []
                    while j + 2 <= len(chunk):
                        code = int.from_bytes(chunk[j : j + 2], "little")
                        if code in _HWP_EXTENDED_CTRL:
                            j += 16  # 확장 컨트롤은 8워드를 차지
                        elif code < 32:
                            if code in (10, 13):
                                buf.append("\n")
                            j += 2
                        else:
                            buf.append(chr(code))
                            j += 2
                    text = "".join(buf).strip()
                    if text:
                        out.append(text)
                i += size
        joined = "\n".join(out)
        # UTF-16 서로게이트 쌍(BMP 밖 문자) 복원
        return joined.encode("utf-16", "surrogatepass").decode("utf-16", "ignore")
    finally:
        ole.close()


def _extract_hwpx(path: str) -> str:
    """HWPX(zip)의 Contents/section*.xml에서 문단 단위로 텍스트 추출. 표 셀 포함."""
    import xml.etree.ElementTree as ET

    parts: list[str] = []
    with zipfile.ZipFile(path) as zf:
        names = sorted(n for n in zf.namelist() if re.fullmatch(r"Contents/section\d+\.xml", n))
        if not names:
            raise ValueError("Contents/section*.xml이 없어 HWPX 형식이 아닙니다")
        for name in names:
            root = ET.fromstring(zf.read(name))
            for el in root.iter():  # 문서 순서 순회 — 표 셀 문단도 제자리에서 등장
                tag = el.tag.rsplit("}", 1)[-1]
                if tag == "p":
                    parts.append("\n")
                elif tag == "t" and el.text:
                    parts.append(el.text)
    return "".join(parts).strip()


def _extract_pdf(path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception:
            raise ValueError("암호화된 PDF는 지원하지 않습니다")
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx(path: str) -> str:
    """DOCX 문단·표를 문서 순서대로 추출. 표는 행 단위로 셀을 ' | '로 잇는다."""
    import docx
    from docx.table import Table

    doc = docx.Document(path)
    parts: list[str] = []

    def add_table(tb: Table) -> None:
        for row in tb.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))

    if hasattr(doc, "iter_inner_content"):
        for item in doc.iter_inner_content():
            if isinstance(item, Table):
                add_table(item)
            else:
                parts.append(item.text)
    else:  # 구버전 python-docx: 순서 보존 불가, 내용은 모두 수집
        parts.extend(p.text for p in doc.paragraphs)
        for tb in doc.tables:
            add_table(tb)
    return "\n".join(parts)


_DOC_EXTRACTORS = {
    "hwp": _extract_hwp,
    "hwpx": _extract_hwpx,
    "pdf": _extract_pdf,
    "docx": _extract_docx,
}


@mcp.tool()
async def fetch_job_document(file_url: str) -> str:
    """채용공고 첨부파일(HWP/HWPX/PDF/DOCX)을 내려받아 본문 텍스트를 추출한다.

    직무기술서·공고문 등 표 중심 문서의 표 안 텍스트도 함께 추출된다.

    Args:
        file_url: 첨부파일 다운로드 URL
            (get_job_files 또는 get_alio_job_detail 결과의 첨부파일 URL)
    """
    try:
        path, filename = await _download_doc(_normalize_doc_url(file_url))
    except httpx.TimeoutException:
        return "오류: 파일 다운로드가 30초를 초과했습니다."
    except httpx.HTTPError as e:
        return f"오류: 파일 다운로드 실패 - {e}"
    except ValueError as e:
        return f"오류: {e}"

    try:
        with open(path, "rb") as fh:
            head = fh.read(16)
        if head.lstrip()[:5].lower() in (b"<!doc", b"<html"):
            return "오류: 다운로드 URL이 파일 대신 웹페이지(HTML)를 반환했습니다. URL을 확인하세요."
        fmt = _detect_doc_format(path, filename)
        if fmt is None:
            return (
                f"오류: 지원하지 않는 파일 형식입니다 (파일명: {filename or '알 수 없음'}). "
                "지원 형식: .hwp .hwpx .pdf .docx"
            )
        try:
            text = _DOC_EXTRACTORS[fmt](path)
        except ValueError as e:
            return f"오류: 텍스트 추출 불가 - {e}"
        except Exception as e:
            return f"오류: 텍스트 추출 실패 ({fmt}) - {type(e).__name__}: {e}"

        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        if not text:
            return f"오류: 텍스트를 추출했지만 내용이 비어 있습니다 (형식: {fmt}, 이미지 스캔 문서일 수 있음)"

        header = f"[문서: {filename or file_url} | 형식: {fmt} | {len(text):,}자]"
        if len(text) > DOC_MAX_CHARS:
            cut = len(text) - DOC_MAX_CHARS
            text = text[:DOC_MAX_CHARS] + f"\n\n[주의: 문서가 길어 앞 {DOC_MAX_CHARS:,}자만 표시. 이후 {cut:,}자 절단됨]"
        return f"{header}\n{text}"
    finally:
        try:
            os.remove(path)
        except OSError:
            pass


# ---------------------------------------------------------------------------
# NCS 국가직무능력표준 (apis.data.go.kr/B490007)
# ---------------------------------------------------------------------------

NCS_BASE_URL = "https://apis.data.go.kr/B490007/ncsInfo"
NCS_JM_URL = "https://apis.data.go.kr/B490007/ncsClCdJm/getNcsClCdJmList"
NCS_CACHE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ncs_cache.json")


def _ncs_key() -> str | None:
    key = os.environ.get("PUBJOB_API_KEY")  # 같은 data.go.kr 계정 키 재사용
    if key and "%" in key:
        key = unquote(key)
    return key


async def _call_ncs(path: str, params: dict[str, Any]) -> dict[str, Any] | str:
    """ncsInfo API 호출. 성공 시 전체 JSON(dict), 실패 시 에러 메시지(str) 반환.

    명세와 달리 returnType=json이 사실상 필수다 (없으면 code=009).
    """
    key = _ncs_key()
    if not key:
        return "오류: 환경변수 PUBJOB_API_KEY가 비어 있습니다."
    query = {"serviceKey": key, "returnType": "json", **params}
    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            resp = await client.get(f"{NCS_BASE_URL}{path}", params=query)
            resp.raise_for_status()
        data = resp.json()
    except httpx.TimeoutException:
        return "오류: NCS API 응답이 10초를 초과했습니다."
    except httpx.HTTPError as e:
        return f"오류: NCS API 요청 실패 - {e}"
    except ValueError:
        return f"오류: NCS API JSON 파싱 실패 / 응답 일부: {resp.text[:200]}"
    info = data.get("dataInfo") or {}
    if info.get("code") not in (None, "000"):
        return f"오류: NCS API 실패 응답 (code={info.get('code')}, message={info.get('message')})"
    return data


def _ncs_load_cache() -> dict[str, Any]:
    try:
        with open(NCS_CACHE_PATH, encoding="utf-8") as f:
            return json.load(f)
    except (OSError, ValueError):
        return {}


def _ncs_save_cache(cache: dict[str, Any]) -> None:
    with open(NCS_CACHE_PATH, "w", encoding="utf-8") as f:
        json.dump(cache, f, ensure_ascii=False)


async def _ncs_fetch_all(path: str, extract: Any) -> tuple[list[Any], int] | str:
    """엔드포인트 전체를 페이지 순회로 수집한다. (rows, 호출수) 또는 에러 메시지."""
    rows: list[Any] = []
    calls = 0
    num_rows = 1000
    page = 1
    while True:
        data = await _call_ncs(path, {"numOfRows": str(num_rows), "pageNo": str(page)})
        calls += 1
        if isinstance(data, str):
            if num_rows == 1000 and page == 1:  # 1000행이 거부되면 100행으로 재시도
                num_rows = 100
                continue
            return data
        batch = data.get("data") or []
        rows.extend(extract(r) for r in batch)
        total = int((data.get("dataInfo") or {}).get("totCnt") or 0)
        if not batch or len(rows) >= total or page > 200:
            return rows, calls
        page += 1


async def _ncs_ensure_cache() -> dict[str, Any] | str:
    """분류(ncsCdInfo)·능력단위(ncsCompeUnitInfo) 전체를 캐시. 있으면 그대로 반환."""
    cache = _ncs_load_cache()
    if cache.get("classifications") and cache.get("units"):
        return cache

    cls_result = await _ncs_fetch_all(
        "/ncsCdInfo",
        lambda r: {
            "dutyCd": f"{r.get('ncsLclasCd')}{r.get('ncsMclasCd')}{r.get('ncsSclasCd')}{r.get('ncsSubdCd')}",
            "path": f"{r.get('ncsLclasCdNm')} > {r.get('ncsMclasCdNm')} > {r.get('ncsSclasCdNm')} > {r.get('ncsSubdCdNm')}",
            "names": [r.get("ncsLclasCdNm"), r.get("ncsMclasCdNm"), r.get("ncsSclasCdNm"), r.get("ncsSubdCdNm")],
        },
    )
    if isinstance(cls_result, str):
        return cls_result
    classifications, cls_calls = cls_result

    unit_result = await _ncs_fetch_all(
        "/ncsCompeUnitInfo",
        lambda r: {
            "dutyCd": r.get("dutyCd"),
            "compUnitCd": r.get("compUnitCd"),
            "name": re.sub(r"^\d+\.", "", r.get("compUnitName") or ""),
            "level": r.get("compUnitLevel"),
            "ncsClCd": r.get("ncsClCd"),
        },
    )
    if isinstance(unit_result, str):
        return unit_result
    unit_rows, unit_calls = unit_result

    units: dict[str, list[dict[str, Any]]] = {}
    for u in unit_rows:
        units.setdefault(u.pop("dutyCd") or "", []).append(u)

    cache = {
        "meta": {
            "built_at": datetime.now().isoformat(timespec="seconds"),
            "calls_used": cls_calls + unit_calls,
            "classification_count": len(classifications),
            "unit_count": len(unit_rows),
        },
        "classifications": classifications,
        "units": units,
        "related_quals": cache.get("related_quals", {}),
    }
    _ncs_save_cache(cache)
    return cache


@mcp.tool()
async def ncs_find_duty(keyword: str) -> str:
    """키워드로 NCS 세분류(직무)를 찾는다. 대/중/소/세분류명과 능력단위명에서 부분일치 검색.

    최초 호출 시 전체 분류·능력단위를 내려받아 로컬 캐시(ncs_cache.json)를 만든다.

    Args:
        keyword: 검색 키워드. 쉼표로 구분하면 OR 검색 (예: "프로젝트관리,사업기획")
    """
    cache = await _ncs_ensure_cache()
    if isinstance(cache, str):
        return cache

    keywords = [k.strip() for k in keyword.split(",") if k.strip()]
    if not keywords:
        return "오류: 검색 키워드를 입력하세요."

    matched: list[tuple[dict[str, Any], list[dict[str, Any]]]] = []
    for c in cache["classifications"]:
        units = cache["units"].get(c["dutyCd"], [])
        name_hit = any(kw in (n or "") for kw in keywords for n in c["names"])
        unit_hit = any(kw in (u["name"] or "") for kw in keywords for u in units)
        if name_hit or unit_hit:
            matched.append((c, units))

    if not matched:
        return f"검색 결과 없음 (키워드: {keyword})"

    lines = [f"[NCS 직무 검색] 키워드: {keyword} / {len(matched)}개 세분류 매칭"]
    shown = matched[:15]
    for c, units in shown:
        lines.append(f"\n■ {c['path']} (dutyCd={c['dutyCd']})")
        for u in units:
            lines.append(f"  - [{u['compUnitCd']}] {u['name']} (수준 {u['level']}, ncsClCd={u['ncsClCd']})")
        if not units:
            lines.append("  (능력단위 정보 없음)")
    if len(matched) > len(shown):
        lines.append(f"\n… 외 {len(matched) - len(shown)}개 세분류 생략. 키워드를 좁혀서 다시 검색하세요.")
    return "\n".join(lines)


@mcp.tool()
async def ncs_duty_overview(duty_cd: str) -> str:
    """NCS 세분류(직무)의 직무정의와 능력단위 목록을 조회한다.

    Args:
        duty_cd: 8자리 세분류 코드 (예: "01010102", ncs_find_duty 결과의 dutyCd)
    """
    duty_data, unit_data = await asyncio.gather(
        _call_ncs("/ncsDutyInfo", {"dutyCd": duty_cd, "numOfRows": "10", "pageNo": "1"}),
        _call_ncs("/ncsCompeUnitInfo", {"dutyCd": duty_cd, "numOfRows": "100", "pageNo": "1"}),
    )

    lines = [f"[NCS 직무 개요] dutyCd={duty_cd}"]
    if isinstance(duty_data, str):
        lines.append(duty_data)
    else:
        for d in duty_data.get("data") or []:
            lines.append(f"직무명: {d.get('dutyNm')}")
            lines.append(f"직무정의: {d.get('dutyDef')}")

    lines.append("\n--- 능력단위 목록 ---")
    if isinstance(unit_data, str):
        lines.append(unit_data)
    else:
        units = unit_data.get("data") or []
        if not units:
            lines.append(f"능력단위 없음 (dutyCd={duty_cd} 확인 필요)")
        for u in units:
            lines.append(
                f"[{u.get('compUnitCd')}] {u.get('compUnitName')} "
                f"(수준 {u.get('compUnitLevel')}, ncsClCd={u.get('ncsClCd')})"
            )
            if u.get("compUnitDef"):
                lines.append(f"    {u['compUnitDef']}")
    return "\n".join(lines)


@mcp.tool()
async def ncs_analyze_unit(duty_cd: str, comp_unit_cd: str, include_exam: bool = False) -> str:
    """NCS 능력단위 1개를 분석한다: 능력단위요소별 수행준거·지식·기술·태도 + 직업기초능력.

    자기소개서·면접 준비, 직무기술서 해석에 필요한 상세 기준을 구조화해 반환한다.

    Args:
        duty_cd: 8자리 세분류 코드 (예: "01010102")
        comp_unit_cd: 2자리 능력단위 코드 (예: "01")
        include_exam: True면 출제기준(평가방법·시험시간·장비)과 평가지침도 포함
    """
    unit = {"dutyCd": duty_cd, "compUnitCd": comp_unit_cd, "pageNo": "1"}
    tasks = [
        _call_ncs("/ncsCompeUnitFactrInfo", {**unit, "numOfRows": "100"}),
        _call_ncs("/ncsKsaInfo", {**unit, "numOfRows": "1000"}),
        _call_ncs("/ncsjobInfo", {**unit, "numOfRows": "1000"}),
    ]
    if include_exam:
        tasks += [
            _call_ncs("/ncsSetqInfo", {**unit, "numOfRows": "100"}),
            _call_ncs("/ncsEvalInfo", {**unit, "numOfRows": "100"}),
        ]
    results = await asyncio.gather(*tasks)
    factr_data, ksa_data, job_data = results[0], results[1], results[2]

    lines = [f"[NCS 능력단위 분석] dutyCd={duty_cd} / compUnitCd={comp_unit_cd}"]

    factors: list[dict[str, Any]] = []
    if isinstance(factr_data, str):
        lines.append(f"능력단위요소: {factr_data}")
    else:
        factors = factr_data.get("data") or []
        if factors:
            f0 = factors[0]
            lines.append(f"능력단위: {f0.get('compUnitName')} (수준 {f0.get('compUnitLevel')}, ncsClCd={f0.get('ncsClCd')})")

    # KSA 행을 (요소번호, 구분명)으로 묶는다. gbnName: 수행준거/지식/기술/태도
    ksa_by_factor: dict[Any, dict[str, list[str]]] = {}
    if isinstance(ksa_data, str):
        lines.append(f"수행준거·KSA: {ksa_data}")
    else:
        for row in ksa_data.get("data") or []:
            grp = ksa_by_factor.setdefault(row.get("compUnitFactrNo"), {})
            grp.setdefault(row.get("gbnName") or "기타", []).append(row.get("gbnVal") or "")

    for f in factors:
        no = f.get("compUnitFactrNo")
        lines.append(f"\n■ 요소 {f.get('compUnitFactrName')}")
        grp = ksa_by_factor.get(no, {})
        for gbn in ("수행준거", "지식", "기술", "태도"):
            vals = grp.pop(gbn, [])
            if vals:
                lines.append(f"  [{gbn}]")
                lines.extend(f"  - {v}" for v in vals)
        for gbn, vals in grp.items():  # 예상 밖 구분값도 누락 없이 표시
            lines.append(f"  [{gbn}]")
            lines.extend(f"  - {v}" for v in vals)

    lines.append("\n--- 직업기초능력 ---")
    if isinstance(job_data, str):
        lines.append(job_data)
    else:
        mains: dict[str, list[str]] = {}
        for row in job_data.get("data") or []:
            mains.setdefault(row.get("mainName") or "", []).append(row.get("subName") or "")
        if not mains:
            lines.append("정보 없음")
        for main, subs in mains.items():
            lines.append(f"- {main}: {', '.join(subs)}")

    if include_exam:
        setq_data, eval_data = results[3], results[4]
        lines.append("\n--- 출제기준 ---")
        if isinstance(setq_data, str):
            lines.append(setq_data)
        else:
            d = setq_data.get("data") or {}
            for e in d.get("evalData") or []:
                lines.append(f"- 평가방법: {e.get('evalMethDstinName')} / {e.get('evalMethName')}")
            for t in d.get("timeData") or []:
                lines.append(f"- 시험시간: 지필 {t.get('papenEvalTime')}분, 실무 {t.get('pracbizEvalTime')}분")
            equips = [e.get("equipName") for e in d.get("equipData") or [] if e.get("equipName")]
            if equips:
                lines.append(f"- 시설·장비: {', '.join(equips)}")
        lines.append("\n--- 평가지침 ---")
        if isinstance(eval_data, str):
            lines.append(eval_data)
        else:
            d = eval_data.get("data") or {}
            seen: set[tuple[Any, Any]] = set()
            for e in d.get("evalData") or []:
                key = (e.get("evalTypeName"), e.get("evalMethName"))
                if key in seen:
                    continue
                seen.add(key)
                desc = f" — {e.get('evalDesc')}" if e.get("evalDesc") else ""
                lines.append(f"- {e.get('evalTypeName')} / {e.get('evalMethName')}{desc}")
    return "\n".join(lines)


@mcp.tool()
async def ncs_related_quals(ncs_cl_cd: str) -> str:
    """NCS 능력단위와 연계된 국가기술자격 종목 목록을 조회한다. 결과는 캐시된다(일 1000건 제한).

    Args:
        ncs_cl_cd: 능력단위 풀코드+버전 (형식: 세분류8자리+능력단위2자리_차수v버전,
            예: "0101010201_17v2" — ncs_find_duty/ncs_duty_overview 결과의 ncsClCd 값 그대로)
    """
    if not re.fullmatch(r"\d{10}_\d+v\d+", ncs_cl_cd):
        return (
            f"오류: ncs_cl_cd 형식이 잘못되었습니다 ({ncs_cl_cd}). "
            "'0101010201_17v2' 형식이어야 합니다 (ncsClCd 값을 그대로 사용)."
        )

    cache = _ncs_load_cache()
    items = cache.get("related_quals", {}).get(ncs_cl_cd)
    from_cache = items is not None

    if items is None:
        key = _ncs_key()
        if not key:
            return "오류: 환경변수 PUBJOB_API_KEY가 비어 있습니다."
        items = []
        page = 1
        while True:  # 페이지당 최대 50건 제한(resultCode=930)이 있어 순회 수집
            try:
                async with httpx.AsyncClient(timeout=TIMEOUT) as client:
                    resp = await client.get(
                        NCS_JM_URL,
                        params={
                            "serviceKey": key,
                            "dataFormat": "json",
                            "numOfRows": "50",
                            "pageNo": str(page),
                            "ncsClCd": ncs_cl_cd,
                        },
                    )
                    resp.raise_for_status()
                data = resp.json()
            except httpx.TimeoutException:
                return "오류: NCS 자격 API 응답이 10초를 초과했습니다."
            except httpx.HTTPError as e:
                return f"오류: NCS 자격 API 요청 실패 - {e}"
            except ValueError:
                return f"오류: NCS 자격 API JSON 파싱 실패 / 응답 일부: {resp.text[:200]}"

            header = data.get("header") or {}
            if header.get("resultCode") != "00":
                return f"오류: NCS 자격 API 실패 응답 (resultCode={header.get('resultCode')}, resultMsg={header.get('resultMsg')})"
            body = data.get("body") or {}
            batch = body.get("items") or []
            if isinstance(batch, dict):
                batch = [batch]
            items.extend(batch)
            total = int(body.get("totalCount") or 0)
            if not batch or len(items) >= total or page > 40:
                break
            page += 1
        cache.setdefault("related_quals", {})[ncs_cl_cd] = items
        _ncs_save_cache(cache)

    if not items:
        return f"연계 자격 없음 (ncsClCd={ncs_cl_cd})"

    lines = [f"[연계 국가기술자격] ncsClCd={ncs_cl_cd} / {len(items)}건" + (" (캐시)" if from_cache else "")]
    for it in items:
        lines.append(
            f"- [{it.get('jmCd')}] {it.get('jmNm')} | {it.get('abltUnitTypNm')} | "
            f"시험기관: {it.get('examInstiNm')} | 능력단위 표준시간: {it.get('minEduTrngTm')}h"
        )
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# 고용24 국민내일배움카드 훈련과정 (work24.go.kr)
# ---------------------------------------------------------------------------

WORK24_LIST_URL = "https://www.work24.go.kr/cm/openApi/call/hr/callOpenApiSvcInfo310L01.do"

WORK24_AREA = {
    "11": "서울", "26": "부산", "27": "대구", "28": "인천", "29": "광주", "30": "대전",
    "31": "울산", "36": "세종", "41": "경기", "42": "강원", "43": "충북", "44": "충남",
    "45": "전북", "46": "전남", "47": "경북", "48": "경남", "50": "제주",
}


def _w24(row: dict[str, Any], *names: str) -> Any:
    """응답 필드명이 명세와 다를 수 있어 후보명을 순서대로 시도한다."""
    for n in names:
        if row.get(n) not in (None, ""):
            return row[n]
    return None


def _ncs_param(code: str) -> dict[str, str]:
    """NCS 코드를 srchNcs1 하나에 그대로 넣는다.

    실측 결과 srchNcs1은 **가변길이 접두 코드**를 받는다 (2/4/6/8자리 모두 동작).
    srchNcs2~4는 분류 단계가 아니라 각자 독립적인 대분류 필터로 동작하며 뒤에 온 값이
    앞을 덮어쓰므로 사용하지 않는다. 자세한 근거는 README quirk 참조.
    """
    code = re.sub(r"\D", "", code or "")[:8]
    return {"srchNcs1": code} if code else {}


async def _call_work24(params: dict[str, Any]) -> dict[str, Any] | str:
    """고용24 훈련과정 API 호출. 성공 시 JSON(dict), 실패 시 에러 메시지(str)."""
    key = os.environ.get("WORK24_API_KEY")
    if not key:
        return "오류: 환경변수 WORK24_API_KEY가 비어 있습니다. .env 파일에 인증키를 입력하세요."
    if "%" in key:
        key = unquote(key)

    query: dict[str, Any] = {"authKey": key, "returnType": "JSON", "outType": "1"}
    query.update({k: v for k, v in params.items() if v not in (None, "")})

    try:
        async with httpx.AsyncClient(timeout=TIMEOUT) as client:
            resp = await client.get(WORK24_LIST_URL, params=query)
            resp.raise_for_status()
    except httpx.TimeoutException:
        return "오류: 고용24 API 응답이 10초를 초과했습니다."
    except httpx.HTTPError as e:
        return f"오류: 고용24 API 요청 실패 - {e}"

    body = resp.text.strip()
    if body[:1] not in ("{", "["):  # 인증 실패 시 XML/HTML 에러를 돌려주는 경우
        return f"오류: JSON이 아닌 응답 (인증키·파라미터 확인) / 응답 일부: {body[:300]}"
    try:
        data = resp.json()
    except ValueError as e:
        return f"오류: JSON 파싱 실패 - {e} / 응답 일부: {body[:200]}"

    if isinstance(data, dict):
        # 일부 오픈API는 결과를 문자열 JSON으로 한 번 더 감싼다
        inner = data.get("returnJSON")
        if isinstance(inner, str):
            try:
                data = json.loads(inner)
            except ValueError:
                pass
        msg = _w24(data, "errMsg", "message", "resultMsg")
        code = str(_w24(data, "errCd", "resultCode") or "")
        if msg and code not in ("", "0", "00", "000"):
            return f"오류: 고용24 API 실패 응답 (code={code}, msg={msg})"
    return data if isinstance(data, dict) else {"srchList": data}


def _w24_rows(data: dict[str, Any]) -> list[dict[str, Any]]:
    for key in ("srchList", "list", "items", "returnList"):
        v = data.get(key)
        if isinstance(v, list):
            return [r for r in v if isinstance(r, dict)]
        if isinstance(v, dict):
            return [v]
    return []


@mcp.tool()
async def search_training_courses(
    keyword: str | None = None,
    ncs_code: str | None = None,
    online_only: bool = False,
    course_type: str | None = None,
    area: str | None = None,
    org_name: str | None = None,
    start_from: str | None = None,
    start_to: str | None = None,
    page_size: int = 20,
) -> str:
    """고용24(국민내일배움카드) 훈련과정을 검색한다. 취업률·만족도·실훈련비까지 반환한다.

    Args:
        keyword: 과정명 검색어 (예: "회계실무", "연구행정")
        ncs_code: NCS 직무코드. **ncs_find_duty/ncs_duty_overview의 dutyCd를 그대로** 넣는다.
            고용24의 ncsCd는 우리 dutyCd와 같은 8자리 체계라 코드가 그대로 통한다.
            접두 검색이라 자릿수로 범위 조절 — "01"=사업관리 대분류 전체,
            "0101"=중분류, "01010102"=프로젝트관리 세분류만.
        online_only: True면 인터넷(원격) 과정만
        course_type: 과정 구분 코드 — C0061S=구직자, C0061I=재직자,
            C0104=K디지털트레이닝, C0105=K디지털기초역량, C0055=실업자원격
        area: 지역 코드 — 11서울 26부산 27대구 28인천 29광주 30대전 31울산 36세종
            41경기 42강원 43충북 44충남 45전북 46전남 47경북 48경남 50제주
        start_from: 훈련 시작일 범위 시작 YYYYMMDD (기본 오늘)
        start_to: 훈련 시작일 범위 끝 YYYYMMDD (기본 오늘+90일)
        page_size: 결과 수 (기본 20, 최대 100)
    """
    today = date.today()
    st = start_from or today.strftime("%Y%m%d")
    ed = start_to or (today + timedelta(days=90)).strftime("%Y%m%d")
    rows_n = max(1, min(page_size, 100))

    params: dict[str, Any] = {
        "srchTraStDt": st,
        "srchTraEndDt": ed,
        "sort": "ASC",
        "sortCol": "2",
        "pageNum": "1",
        "pageSize": str(rows_n),
        "srchTraProcessNm": keyword,
        "srchTraOrganNm": org_name,
        "crseTracseSe": course_type,
        "srchTraArea1": area,
    }
    if online_only:
        params["srchTraGbn"] = "M1005"
    if ncs_code:
        params.update(_ncs_param(ncs_code))

    data = await _call_work24(params)
    if isinstance(data, str):
        return data
    rows = _w24_rows(data)

    # NCS 4단 분절 가설이 틀려 0건이면 대분류(srchNcs1)만으로 폴백한다
    fallback_note = ""
    if not rows and ncs_code and len(re.sub(r"\D", "", ncs_code)) > 2:
        p2 = dict(params)
        for i in (2, 3, 4):
            p2.pop(f"srchNcs{i}", None)
        data2 = await _call_work24(p2)
        if not isinstance(data2, str):
            rows2 = _w24_rows(data2)
            if rows2:
                data, rows = data2, rows2
                fallback_note = (
                    "\n※ NCS 4단 분절(srchNcs1~4) 검색은 0건이라 대분류(srchNcs1)만으로 재검색했습니다."
                )

    conds = [f"훈련시작일: {st}~{ed}"]
    if keyword:
        conds.append(f"과정명: {keyword}")
    if ncs_code:
        conds.append(f"NCS: {_ncs_param(ncs_code).get('srchNcs1', ncs_code)}")
    if online_only:
        conds.append("인터넷과정만")
    if course_type:
        conds.append(f"과정구분: {course_type}")
    if area:
        conds.append(f"지역: {WORK24_AREA.get(area, area)}")
    if org_name:
        conds.append(f"기관: {org_name}")
    cond_text = " / ".join(conds)

    if not rows:
        return f"검색 결과 없음\n[검색 조건] {cond_text}{fallback_note}"

    total = _w24(data, "scn_cnt", "totalCount", "cnt")
    lines = [f"[검색 조건] {cond_text}"]
    if total is not None:
        lines.append(f"전체 {total}건 중 {len(rows)}건 표시")
    if fallback_note:
        lines.append(fallback_note.strip())

    for r in rows:
        title = _fmt_value(_w24(r, "title", "trprNm", "traProcessNm"))
        org = _fmt_value(_w24(r, "subTitle", "trainstCstNm", "traOrganNm"))
        addr = _fmt_value(_w24(r, "address", "addr", "trngAreaNm"))
        sdt = _fmt_value(_w24(r, "traStartDate", "trStartDate", "traStDt"))
        edt = _fmt_value(_w24(r, "traEndDate", "trEndDate", "traEndDt"))
        gbn = _w24(r, "trainTarget", "trprGbn", "srchTraGbnNm")
        ncs = _fmt_value(_w24(r, "ncsCd", "ncsCdNm", "ncsNm"))
        cost = _w24(r, "courseMan", "realMan")
        real = _w24(r, "realMan", "courseMan")
        emp3 = _w24(r, "eiEmplRate3", "eiEmplCnt3Rate", "employ3Rate")
        emp6 = _w24(r, "eiEmplRate6", "eiEmplCnt6Rate", "employ6Rate")
        grade = _w24(r, "grade", "stdgScor", "satisfaction")
        yard = _w24(r, "yardMan", "trainTargetCnt")
        reg = _w24(r, "regCourseMan", "regCnt")
        link = _w24(r, "titleLink", "trprLink")
        tid, tdeg = _w24(r, "trprId"), _w24(r, "trprDegr")

        lines.append(f"\n■ {title}")
        lines.append(f"  기관: {org} | 지역: {addr} | 기간: {sdt}~{edt}")
        extra = []
        if gbn:
            extra.append(f"구분: {_fmt_value(gbn)}")
        if ncs != "-":
            extra.append(f"NCS: {ncs}")
        if cost is not None:
            extra.append(f"수강비: {_fmt_value(cost)}원")
        if real is not None:
            extra.append(f"실훈련비: {_fmt_value(real)}원")
        if extra:
            lines.append("  " + " | ".join(extra))
        perf = []
        if emp3 is not None:
            perf.append(f"취업률(3개월): {_fmt_value(emp3)}%")
        if emp6 is not None:
            perf.append(f"취업률(6개월): {_fmt_value(emp6)}%")
        if grade is not None:
            perf.append(f"만족도: {_fmt_value(grade)}")
        if yard is not None or reg is not None:
            perf.append(f"정원/신청: {_fmt_value(yard)}/{_fmt_value(reg)}")
        if perf:
            lines.append("  " + " | ".join(perf))
        if link:
            lines.append(f"  링크: {_fmt_value(link)}")
        if tid:
            lines.append(f"  (상세조회용) trprId={_fmt_value(tid)} trprDegr={_fmt_value(tdeg)}")
    return "\n".join(lines)


if __name__ == "__main__":
    mcp.run()  # stdio
